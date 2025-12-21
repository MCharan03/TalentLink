import streamlit as st
import base64
import io
import os
from pdfminer.layout import LAParams, LTTextBox
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.converter import TextConverter
import docx
import ast
from fpdf import FPDF
import textwrap
import datetime
import pandas as pd
import pytesseract
from pdf2image import convert_from_path
import yt_dlp


def format_as_bullet_list(text, bullet_char="*"):
    try:
        items = ast.literal_eval(text)
        if isinstance(items, list):
            return "\n".join([f"{bullet_char} {item.strip()}" for item in items])
    except (ValueError, SyntaxError):
        return text
    return text


def fetch_yt_video_title(link):
    """Fetches YouTube video title using yt-dlp, with robust error handling."""
    try:
        ydl_opts = {'quiet': True, 'extract_flat': 'in_playlist',
                    'ignoreerrors': True, 'no_warnings': True}
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:  # type: ignore
            info = ydl.extract_info(link, download=False)
            return info.get('title', "Video Title Not Available")
    except Exception:
        return "Video Title Not Available"


def get_pdf_text(file_path):
    """Extracts text and page count from a PDF file, with OCR fallback."""
    print(f"Processing file: {file_path}")
    page_count = 0
    # Try extracting text with pdfminer
    resource_manager = PDFResourceManager()
    fake_file_handle = io.StringIO()
    converter = TextConverter(
        resource_manager, fake_file_handle, laparams=LAParams())
    page_interpreter = PDFPageInterpreter(resource_manager, converter)
    with open(file_path, 'rb') as fh:
        for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
            page_interpreter.process_page(page)
            page_count += 1
        text = fake_file_handle.getvalue()
    converter.close()
    fake_file_handle.close()

    print(f"Extracted text length: {len(text)}")
    # If text is short, try OCR
    if len(text.strip()) < 100:
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(file_path)
            page_count = len(images)
            ocr_text = []
            for image in images:
                ocr_text.append(pytesseract.image_to_string(image))
            text = "\n".join(ocr_text)
            print(f"OCR extracted text length: {len(text)}")
        except ImportError:
            print("pdf2image is not installed. Skipping OCR.")
        except Exception as e:
            print(f"Error during OCR: {e}")
            st.error(f"Error during OCR: {e}")
            st.info(
                "Please make sure you have Tesseract and Poppler installed and configured correctly.")

    print(f"Returning text length: {len(text)}, page count: {page_count}")
    return text, page_count


def create_user_pdf(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    font_name = "Helvetica"
    bullet_char = "*"

    # Header
    pdf.set_font(font_name, 'B', 20)
    pdf.set_text_color(40, 40, 40)
    pdf.cell(0, 20, 'Resume Analysis Report', 0, align='C', ln=1)
    pdf.set_font(font_name, '', 10)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(
        0, 10, f"Report generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, align='C', ln=1)
    pdf.ln(15)

    pdf.set_text_color(0, 0, 0)

    def render_section(title, data):
        pdf.set_font(font_name, 'B', 14)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(0, 10, title, 0, align='L', fill=True, ln=1)
        pdf.ln(5)
        for key, value in data.items():
            pdf.set_font(font_name, 'B', 12)
            pdf.cell(50, 8, f"{key}:")
            pdf.set_font(font_name, '', 12)
            formatted_value = str(value)
            if key in ["Skills", "Recommended Skills", "Recommended Courses"]:
                formatted_value = format_as_bullet_list(
                    str(value), bullet_char=bullet_char)
            pdf.multi_cell(0, 8, formatted_value.encode(
                'latin-1', 'replace').decode('latin-1'))
            pdf.ln(3)
        pdf.ln(10)

    user_data = df.iloc[0]

    personal_info = {
        "Name": user_data.get("Name", "N/A"),
        "Email ID": user_data.get("Email_ID", "N/A"),
    }
    render_section("Personal Information", personal_info)

    analysis_info = {
        "Predicted Field": user_data.get("Predicted_Field", "N/A"),
        "Experience Level": user_data.get("User_level", "N/A"),
        "AI Feedback": user_data.get('ai_feedback', 'N/A'),
        "Skills": user_data.get("Actual_skills", "N/A"),
    }
    render_section("Resume Analysis", analysis_info)

    recommendations_info = {
        "Recommended Skills": user_data.get("Recommended_skills", "N/A"),
        "Recommended Courses": user_data.get("Recommended_courses", "N/A"),
    }
    render_section("Recommendations", recommendations_info)

    bonus_videos_info = {
        "Resume Writing Tip Video": user_data.get("Resume Writing Tip URL", "N/A"),
        "Interview Tip Video": user_data.get("Interview Tip URL", "N/A"),
    }
    render_section("Bonus Videos", bonus_videos_info)

    # Footer
    pdf.set_y(-15)
    pdf.set_font(font_name, 'I', 8)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 10, f'Page {pdf.page_no()}', 0, align='C', ln=0)

    return pdf.output(dest='S').encode('latin-1')


def create_doc(df):
    document = docx.Document()
    document.add_heading('Resume Analysis Report', 0)
    document.add_paragraph()

    if len(df) == 1:  # User report format
        for col in df.columns:
            document.add_heading(col, level=1)
            value = df.iloc[0][col]
            formatted_value = str(value)
            if col in ["Actual_skills", "Recommended_skills", "Recommended_courses"]:
                formatted_value = format_as_bullet_list(str(value))
            document.add_paragraph(formatted_value)
    else:  # Admin report format (table)
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = col_name

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col_name in enumerate(df.columns):
                row_cells[i].text = str(row[col_name])

    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()


def create_admin_pdf(df):
    pdf = FPDF(orientation='L')
    pdf.add_page()

    font_name = "Helvetica"

    pdf.set_font(font_name, size=6)
    page_width = pdf.w - 2 * pdf.l_margin

    # Define column widths
    col_widths = {
        "ID": 30,
        "Name": 50,
        "Email_ID": 60,
        "ai_feedback": 150,
        "Timestamp": 40,
        "Page_no": 20,
        "Predicted_Field": 50,
        "User_level": 40,
        "Actual_skills": 150,
        "Recommended_skills": 150,
        "Recommended_courses": 150,
    }

    df_cols = [col for col in col_widths if col in df.columns]

    # Adjust column widths to fit the page
    total_width = sum(col_widths[col] for col in df_cols)
    scale_factor = page_width / total_width if total_width > 0 else 1
    for col in df_cols:
        col_widths[col] = max(10, int(col_widths[col] * scale_factor))

    line_height = pdf.font_size * 2.5

    def draw_header():
        pdf.set_font(font_name, 'B', 8)
        pdf.set_fill_color(240, 240, 240)
        for col in df_cols:
            pdf.cell(col_widths[col], line_height, str(col), 1, 0, 'C', True)
        pdf.ln()

    draw_header()

    pdf.set_font(font_name, '', 8)
    pdf.set_fill_color(255, 255, 255)

    for index, row in df.iterrows():
        # Determine the maximum height required for the row
        max_h = line_height
        for col in df_cols:
            text = str(row[col]).encode('latin-1', 'replace').decode('latin-1')
            lines = pdf.multi_cell(
                w=col_widths[col] - 4, h=line_height, txt=text, split_only=True)
            h = (len(lines) * line_height)
            if h > max_h:
                max_h = h

        # Check for page break before drawing the row
        if pdf.get_y() + max_h > pdf.page_break_trigger:
            pdf.add_page()
            draw_header()
            pdf.set_font(font_name, '', 8)

        y_start = pdf.get_y()
        x_start = pdf.get_x()

        for i, col in enumerate(df_cols):
            x_pos = x_start + sum(col_widths[c] for c in df_cols[:i])
            pdf.rect(x_pos, y_start, col_widths[col], max_h)
            pdf.set_xy(x_pos + 2, y_start + 2)
            pdf.multi_cell(col_widths[col] - 4, line_height, str(row[col]
                                                                 ).encode('latin-1', 'replace').decode('latin-1'), align='L')

        pdf.set_y(y_start + max_h)

    return pdf.output(dest='S').encode('latin-1')


def show_pdf(file_path):
    """Displays a PDF file in an iframe."""
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = F'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)
