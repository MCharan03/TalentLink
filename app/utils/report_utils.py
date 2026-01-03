from fpdf import FPDF
from docx import Document
import os
import datetime


def generate_pdf_report(result, file_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 20)
    pdf.set_text_color(40, 40, 40)
    pdf.cell(0, 15, "Resume Analysis Report", 0, 1, 'C')
    
    pdf.set_font("Arial", 'I', 10)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 5, f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, 'C')
    pdf.ln(10)

    # Basic Info
    pdf.set_font("Arial", 'B', 14)
    pdf.set_text_color(0, 0, 0)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(0, 10, "Candidate Information", 0, 1, 'L', True)
    pdf.ln(3)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(40, 8, "Name:")
    pdf.set_font("Arial", '', 11)
    pdf.cell(0, 8, str(result.get('name', 'N/A')), 0, 1)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(40, 8, "Predicted Field:")
    pdf.set_font("Arial", '', 11)
    pdf.cell(0, 8, str(result.get('predicted_field', 'N/A')), 0, 1)

    pdf.set_font("Arial", 'B', 11)
    pdf.cell(40, 8, "Experience Level:")
    pdf.set_font("Arial", '', 11)
    pdf.cell(0, 8, str(result.get('experience_level', 'N/A')), 0, 1)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(40, 8, "Resume Score:")
    pdf.set_font("Arial", 'B', 11)
    pdf.set_text_color(30, 144, 255)
    pdf.cell(0, 8, f"{result.get('resume_score', '0')}/100", 0, 1)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(5)

    # Summary
    pdf.set_font("Arial", 'B', 14)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(0, 10, "AI Summary & Analysis", 0, 1, 'L', True)
    pdf.ln(3)
    pdf.set_font("Arial", '', 10)
    # Removing potential HTML tags from ai_summary if any
    summary = result.get('ai_summary', 'Not available.').replace('<br>', '\n').replace('<strong>', '').replace('</strong>', '').replace('<li>', '- ').replace('</li>', '').replace('<ul>', '').replace('</ul>', '')
    pdf.multi_cell(0, 6, summary.encode('latin-1', 'replace').decode('latin-1'))
    pdf.ln(5)

    # Skills
    pdf.set_font("Arial", 'B', 14)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(0, 10, "Skills Analysis", 0, 1, 'L', True)
    pdf.ln(3)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(0, 8, "Detected Skills:", 0, 1)
    pdf.set_font("Arial", '', 10)
    skills_str = ", ".join(result.get('actual_skills', []))
    pdf.multi_cell(0, 6, skills_str if skills_str else "None detected.")
    pdf.ln(3)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(0, 8, "Recommended Skills:", 0, 1)
    pdf.set_font("Arial", '', 10)
    rec_skills_str = ", ".join(result.get('recommended_skills', []))
    pdf.multi_cell(0, 6, rec_skills_str if rec_skills_str else "No recommendations.")
    pdf.ln(5)

    # Recommendations
    pdf.set_font("Arial", 'B', 14)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(0, 10, "Next Steps", 0, 1, 'L', True)
    pdf.ln(3)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(0, 8, "Recommended Courses:", 0, 1)
    pdf.set_font("Arial", '', 10)
    for course in result.get('recommended_courses', []):
        pdf.cell(0, 6, f"- {course}", 0, 1)
    
    pdf.ln(5)
    pdf.set_font("Arial", 'I', 8)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 10, "End of Report", 0, 0, 'C')

    pdf.output(file_path)


def generate_docx_report(result, file_path):
    doc = Document()
    doc.add_heading('Resume Analysis Report', level=0)
    
    p = doc.add_paragraph()
    p.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    doc.add_heading('Candidate Information', level=1)
    doc.add_paragraph(f"Name: {result.get('name', 'N/A')}")
    doc.add_paragraph(f"Predicted Field: {result.get('predicted_field', 'N/A')}")
    doc.add_paragraph(f"Experience Level: {result.get('experience_level', 'N/A')}")
    doc.add_paragraph(f"Overall Score: {result.get('resume_score', '0')}/100")

    doc.add_heading('AI Summary & Analysis', level=1)
    summary = result.get('ai_summary', 'Not available.').replace('<br>', '\n').replace('<strong>', '').replace('</strong>', '').replace('<li>', '- ').replace('</li>', '').replace('<ul>', '').replace('</ul>', '')
    doc.add_paragraph(summary)

    doc.add_heading('Skills Analysis', level=1)
    doc.add_heading('Detected Skills', level=2)
    doc.add_paragraph(", ".join(result.get('actual_skills', [])) if result.get('actual_skills') else "None detected.")
    
    doc.add_heading('Recommended Skills', level=2)
    doc.add_paragraph(", ".join(result.get('recommended_skills', [])) if result.get('recommended_skills') else "No recommendations.")

    doc.add_heading('Next Steps', level=1)
    doc.add_heading('Recommended Courses', level=2)
    for course in result.get('recommended_courses', []):
        doc.add_paragraph(course, style='List Bullet')

    doc.save(file_path)


def generate_resume_pdf(data, file_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Header
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, data['full_name'], 0, 1, 'C')
    pdf.set_font("Arial", '', 10)
    pdf.cell(
        0, 5, f"{data['email']} | {data['phone']} | {data['linkedin']}", 0, 1, 'C')
    pdf.ln(5)

    # Summary
    if data['summary']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Professional Summary", 0, 1)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 5, data['summary'])
        pdf.ln(5)

    # Education
    if data['education_school']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Education", 0, 1)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 5, data['education_school'], 0, 1)
        pdf.set_font("Arial", '', 11)
        pdf.cell(
            0, 5, f"{data['education_degree']} - {data['education_year']}", 0, 1)
        pdf.ln(5)

    # Experience
    if data['job_title']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Experience", 0, 1)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 5, f"{data['job_title']} at {data['company']}", 0, 1)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 5, data['job_description'])
        pdf.ln(5)

    # Skills
    if data['skills']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Skills", 0, 1)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 5, data['skills'])

    pdf.output(file_path)


def generate_resume_pdf_from_profile(data, file_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    def safe_text(text):
        if not text: return ""
        return text.encode('latin-1', 'replace').decode('latin-1')
    
    # 1. Header (Name & Contact)
    pdf.set_font("Arial", 'B', 24)
    pdf.cell(0, 10, safe_text(data.get('name', 'Resume')), 0, 1, 'C')
    
    pdf.set_font("Arial", '', 12)
    contact_info = []
    if data.get('email'): contact_info.append(safe_text(data.get('email')))
    if data.get('phone'): contact_info.append(safe_text(data.get('phone')))
    # Join with separator
    pdf.cell(0, 6, " | ".join(contact_info), 0, 1, 'C')
    
    # Headline
    if data.get('headline'):
        pdf.ln(2)
        pdf.set_font("Arial", 'I', 11)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 5, safe_text(data.get('headline')), align='C')
        pdf.set_text_color(0, 0, 0)
    pdf.ln(5)

    # 2. Summary (About)
    if data.get('about'):
        pdf.set_font("Arial", 'B', 14)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(0, 8, "SUMMARY", 0, 1, 'L', True)
        pdf.ln(2)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 5, safe_text(data.get('about')))
        pdf.ln(5)

    # 3. Experience
    if data.get('experience'):
        pdf.set_font("Arial", 'B', 14)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(0, 8, "EXPERIENCE", 0, 1, 'L', True)
        pdf.ln(2)
        
        for job in data.get('experience', []):
            pdf.set_font("Arial", 'B', 12)
            # Title & Date
            title = safe_text(job.get('title', 'Job Title'))
            dates = safe_text(job.get('dates', ''))
            pdf.cell(130, 6, title, 0, 0)
            pdf.set_font("Arial", 'I', 10)
            pdf.cell(0, 6, dates, 0, 1, 'R')
            
            # Company & Location
            pdf.set_font("Arial", 'B', 11)
            pdf.set_text_color(80, 80, 80)
            company = safe_text(job.get('company', 'Company'))
            location = safe_text(job.get('location', ''))
            pdf.cell(0, 6, f"{company} - {location}", 0, 1)
            pdf.set_text_color(0, 0, 0)
            
            # Description
            desc = safe_text(job.get('description', ''))
            pdf.set_font("Arial", '', 10)
            pdf.multi_cell(0, 5, desc)
            pdf.ln(4)

    # 4. Education
    if data.get('education'):
        pdf.set_font("Arial", 'B', 14)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(0, 8, "EDUCATION", 0, 1, 'L', True)
        pdf.ln(2)
        
        for edu in data.get('education', []):
            pdf.set_font("Arial", 'B', 12)
            school = safe_text(edu.get('school', 'University'))
            dates = safe_text(edu.get('dates', ''))
            pdf.cell(130, 6, school, 0, 0)
            pdf.set_font("Arial", 'I', 10)
            pdf.cell(0, 6, dates, 0, 1, 'R')
            
            pdf.set_font("Arial", '', 11)
            degree = safe_text(edu.get('degree', ''))
            pdf.cell(0, 6, degree, 0, 1)
            pdf.ln(3)

    # 5. Skills
    if data.get('skills'):
        pdf.set_font("Arial", 'B', 14)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(0, 8, "SKILLS", 0, 1, 'L', True)
        pdf.ln(2)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 5, safe_text(data.get('skills')))
        pdf.ln(5)
        
    # 6. Projects
    if data.get('projects'):
        pdf.set_font("Arial", 'B', 14)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(0, 8, "PROJECTS", 0, 1, 'L', True)
        pdf.ln(2)
        
        for proj in data.get('projects', []):
            pdf.set_font("Arial", 'B', 11)
            pdf.cell(0, 6, safe_text(proj.get('name', 'Project')), 0, 1)
            pdf.set_font("Arial", '', 10)
            pdf.multi_cell(0, 5, safe_text(proj.get('description', '')))
            pdf.ln(3)

    pdf.output(file_path)
