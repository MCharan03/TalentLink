import os
from flask import render_template, redirect, url_for, flash, request, send_file, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from fpdf import FPDF
from docx import Document
from flask_socketio import emit

from . import user
from .forms import ResumeAnalysisForm, MockTestForm
from ..extensions import db, socketio
from ..models import UserData, MockTest, JobPosting, JobApplication, Notification
from ..utils.file_utils import extract_text_from_pdf
from ..utils.ai_utils import analyze_resume, get_interview_question, generate_mock_test


@user.route('/dashboard')
@login_required
def dashboard():
    return render_template('user/dashboard.html')

# --- Resume Analysis ---
@user.route('/resume_analysis', methods=['GET', 'POST'])
@login_required
def resume_analysis():
    form = ResumeAnalysisForm()
    if form.validate_on_submit():
        if form.resume.data:
            resume_file = form.resume.data
            job_description = form.job_description.data
            filename = secure_filename(f"{current_user.id}_{resume_file.filename}")
            upload_folder = current_app.config['UPLOAD_FOLDER']
            if not os.path.exists(upload_folder):
                os.makedirs(upload_folder)
            resume_path = os.path.join(upload_folder, filename)
            resume_file.save(resume_path)

            resume_text = extract_text_from_pdf(resume_path)
            analysis_result = analyze_resume(resume_text, job_description)

            user_data = UserData(user_id=current_user.id,
                                 resume_path=filename,
                                 analysis_result=analysis_result)
            db.session.add(user_data)
            db.session.commit()
            
            flash('Resume analyzed successfully!')
            return redirect(url_for('user.resume_analysis_result', user_data_id=user_data.id))
    return render_template('user/resume_analysis.html', form=form)

@user.route('/resume_analysis_result/<int:user_data_id>')
@login_required
def resume_analysis_result(user_data_id):
    user_data = UserData.query.get_or_404(user_data_id)
    return render_template('user/resume_analysis_result.html', result=user_data.analysis_result, user_data_id=user_data.id)

@user.route('/download_report/<int:user_data_id>/<format>')
@login_required
def download_report(user_data_id, format):
    user_data = UserData.query.get_or_404(user_data_id)
    result = user_data.analysis_result
    report_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"report_{user_data.id}.{format}")

    if format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "Resume Analysis Report", 0, 1, 'C')
        pdf.ln(10)

        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, f"Overall Score: {result.get('score', 'N/A')}/100", 0, 1)
        pdf.ln(5)

        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Summary", 0, 1)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 10, result.get('summary', 'Not available.'))
        pdf.ln(5)

        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Matching Skills", 0, 1)
        pdf.set_font("Arial", '', 12)
        for skill in result.get('matching_skills', []):
            pdf.cell(0, 10, f"- {skill}", 0, 1)
        pdf.ln(5)

        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Missing Skills", 0, 1)
        pdf.set_font("Arial", '', 12)
        for skill in result.get('missing_skills', []):
            pdf.cell(0, 10, f"- {skill}", 0, 1)
        pdf.ln(5)

        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Improvement Suggestions", 0, 1)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 10, result.get('improvement_suggestions', 'Not available.'))
        
        pdf.output(report_path)
        return send_file(report_path, as_attachment=True)

    elif format == 'docx':
        doc = Document()
        doc.add_heading('Resume Analysis Report', level=1)
        doc.add_heading(f"Overall Score: {result.get('score', 'N/A')}/100", level=2)
        
        doc.add_heading('Summary', level=3)
        doc.add_paragraph(result.get('summary', 'Not available.'))
        
        doc.add_heading('Matching Skills', level=3)
        for skill in result.get('matching_skills', []):
            doc.add_paragraph(skill, style='List Bullet')

        doc.add_heading('Missing Skills', level=3)
        for skill in result.get('missing_skills', []):
            doc.add_paragraph(skill, style='List Bullet')

        doc.add_heading('Improvement Suggestions', level=3)
        doc.add_paragraph(result.get('improvement_suggestions', 'Not available.'))
        
        doc.save(report_path)
        return send_file(report_path, as_attachment=True)

    return redirect(url_for('user.resume_analysis_result', user_data_id=user_data.id))

# --- Job Applications ---
@user.route('/jobs')
@login_required
def jobs():
    jobs = JobPosting.query.all()
    return render_template('user/jobs.html', jobs=jobs)

@user.route('/apply/<int:job_id>')
@login_required
def apply(job_id):
    job = JobPosting.query.get_or_404(job_id)
    
    if JobApplication.query.filter_by(user_id=current_user.id, job_id=job.id).first():
        flash('You have already applied for this job.')
        return redirect(url_for('user.jobs'))

    user_data = UserData.query.filter_by(user_id=current_user.id).order_by(UserData.uploaded_at.desc()).first()
    if not user_data:
        flash('Please upload a resume before applying for a job.')
        return redirect(url_for('user.resume_analysis'))

    application = JobApplication(user_id=current_user.id,
                                 job_id=job.id,
                                 resume_path=user_data.resume_path)
    db.session.add(application)
    
    notification = Notification(user_id=job.created_by,
                                message=f'{current_user.username} has applied for your job posting: {job.title}')
    db.session.add(notification)
    
    db.session.commit()
    flash('You have successfully applied for the job.')
    return redirect(url_for('user.jobs'))

# --- Mock Interview ---
@user.route('/mock_interview')
@login_required
def mock_interview():
    return render_template('user/mock_interview.html')

@socketio.on('start_interview')
def handle_start_interview(data):
    question = get_interview_question("Tell me about yourself.", "start")
    emit('interview_question', {'question': question})

@socketio.on('send_answer')
def handle_send_answer(data):
    answer = data['answer']
    question = get_interview_question(answer, "continue")
    emit('interview_question', {'question': question})

# --- Mock Test ---
@user.route('/mock_test', methods=['GET', 'POST'])
@login_required
def mock_test():
    form = MockTestForm()
    if form.validate_on_submit():
        topic = form.topic.data
        test_data = generate_mock_test(topic)
        if test_data:
            mock_test = MockTest(user_id=current_user.id,
                                 questions=test_data)
            db.session.add(mock_test)
            db.session.commit()
            return render_template('user/mock_test_start.html', test=mock_test)
        else:
            flash('Could not generate a test for the given topic.')
    return render_template('user/mock_test.html', form=form)

@user.route('/submit_test/<int:test_id>', methods=['POST'])
@login_required
def submit_test(test_id):
    test = MockTest.query.get_or_404(test_id)
    score = 0
    for i, q in enumerate(test.questions['questions']):
        selected_answer = request.form.get(f'question-{i}')
        if selected_answer == q['correct_answer']:
            score += 1
    
    test.score = (score / len(test.questions['questions'])) * 100
    db.session.commit()

    flash(f'You scored {test.score}% on the test.')
    return redirect(url_for('user.dashboard'))
